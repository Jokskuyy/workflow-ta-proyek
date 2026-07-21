#!/usr/bin/env python3
"""Create a clean Faiz-specific DOCX skeleton from project facts.

The generated document is built from scratch. It does not preserve the cover,
approval sheet, cached table of contents, or body content of another member's
report. The normal report pipeline subsequently replaces the BAB I placeholder
with ``Tugas_Akhir_Draft.md`` and applies the campus formatter.
"""

from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUTPUT = ROOT / "scratch" / "Tugas_Akhir_Faiz_Base.docx"
FACTS_PATH = ROOT / "project_facts.json"
LOGO_PATH = ROOT / "images" / "cover_upn_logo.jpg"


def _font_path(bold: bool = False) -> str:
    filename = "timesbd.ttf" if bold else "times.ttf"
    path = Path("C:/Windows/Fonts") / filename
    if not path.exists():
        raise FileNotFoundError(f"Required Times New Roman font is missing: {path}")
    return str(path)


def _load_metadata() -> dict:
    data = json.loads(FACTS_PATH.read_text(encoding="utf-8"))
    return data["project_metadata"]


def _set_run_font(run, size_pt: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), "Times New Roman")


def _add_centered_paragraph(
    doc: Document,
    text: str = "",
    *,
    size_pt: float = 12,
    bold: bool = False,
    before_pt: float = 0,
    after_pt: float = 0,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    if text:
        _set_run_font(paragraph.add_run(text), size_pt=size_pt, bold=bold)
    return paragraph


def _wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        candidate = " ".join(current + [word])
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if current and width > max_width:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


def _draw_centered_lines(draw, lines, font, page_width, y, gap=16):
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        draw.text(((page_width - width) / 2, y), line, fill="black", font=font)
        y += height + gap
    return y


def _create_approval_image(path: Path, metadata: dict) -> None:
    width, height = 2480, 3508
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(_font_path(True), 58)
    heading_font = ImageFont.truetype(_font_path(True), 46)
    body_font = ImageFont.truetype(_font_path(False), 40)
    body_bold = ImageFont.truetype(_font_path(True), 40)

    y = 220
    y = _draw_centered_lines(draw, ["LEMBAR PENGESAHAN", "TUGAS AKHIR"], title_font, width, y, 28)
    y += 90
    title_lines = _wrap_text(draw, metadata["title"].upper(), heading_font, 1900)
    y = _draw_centered_lines(draw, title_lines, heading_font, width, y, 24)
    y += 110
    y = _draw_centered_lines(draw, ["Disusun oleh:"], body_font, width, y, 20)
    y = _draw_centered_lines(draw, [metadata["author"], metadata["nim"]], body_bold, width, y, 20)
    y += 80
    y = _draw_centered_lines(
        draw,
        [metadata["program_study"], metadata["faculty"], metadata["institution"]],
        body_font,
        width,
        y,
        20,
    )
    y += 140
    y = _draw_centered_lines(draw, ["Telah disetujui oleh:"], body_font, width, y, 28)
    y += 30
    for label, name in (
        ("Pembimbing I", metadata["advisors"]["pembimbing_1"]),
        ("Pembimbing II", metadata["advisors"]["pembimbing_2"]),
    ):
        y = _draw_centered_lines(draw, [label], body_bold, width, y, 18)
        y += 90
        draw.line((690, y, 1790, y), fill="black", width=3)
        y += 22
        y = _draw_centered_lines(draw, [name], body_font, width, y, 36)
        y += 60

    note_font = ImageFont.truetype(_font_path(False), 30)
    _draw_centered_lines(
        draw,
        ["[TBD: tanda tangan dan tanggal pengesahan resmi]"],
        note_font,
        width,
        height - 260,
        0,
    )
    image.save(path, format="PNG", optimize=True)


def _add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "Daftar isi akan diperbarui otomatis."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, display, end))


def create_base_docx(output_path: Path = DEFAULT_OUTPUT) -> Path:
    metadata = _load_metadata()
    if not LOGO_PATH.exists():
        raise FileNotFoundError(f"Official UPNVJ logo is missing: {LOGO_PATH}")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14 if style_name == "Heading 1" else 12)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    if "Table of Figures" not in [style.name for style in doc.styles]:
        tof_style = doc.styles.add_style("Table of Figures", WD_STYLE_TYPE.PARAGRAPH)
        tof_style.base_style = normal

    _add_centered_paragraph(
        doc,
        metadata["title"].upper(),
        size_pt=14,
        bold=True,
        after_pt=10,
    )
    _add_centered_paragraph(doc, "TUGAS AKHIR", size_pt=14, bold=True, after_pt=16)
    logo_p = _add_centered_paragraph(doc, after_pt=14)
    logo_p.add_run().add_picture(str(LOGO_PATH), width=Cm(4.2))
    _add_centered_paragraph(doc, "Disusun oleh:", after_pt=4)
    _add_centered_paragraph(doc, metadata["author"], bold=True, after_pt=2)
    _add_centered_paragraph(doc, metadata["nim"], bold=True, after_pt=14)
    _add_centered_paragraph(doc, metadata["program_study"].upper(), bold=True)
    _add_centered_paragraph(doc, metadata["faculty"].upper(), bold=True)
    _add_centered_paragraph(doc, metadata["institution"].upper(), bold=True)
    _add_centered_paragraph(doc, "2026", bold=True)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    with tempfile.TemporaryDirectory() as tmp_dir:
        approval_path = Path(tmp_dir) / "faiz_approval.png"
        _create_approval_image(approval_path, metadata)
        approval_p = doc.add_paragraph()
        approval_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        approval_p.paragraph_format.first_line_indent = Cm(0)
        approval_p.add_run().add_picture(str(approval_path), width=Cm(14))

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    toc_heading = doc.add_paragraph("DAFTAR ISI", style="Heading 1")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.first_line_indent = Cm(0)
    _add_toc_field(toc_p)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    figures_heading = doc.add_paragraph("DAFTAR GAMBAR", style="Heading 1")
    figures_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("", style="Table of Figures")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    tables_heading = doc.add_paragraph("DAFTAR TABEL", style="Heading 1")
    tables_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("", style="Table of Figures")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("BAB I PENDAHULUAN", style="Heading 1")

    props = doc.core_properties
    props.title = metadata["title"]
    props.subject = "Tugas Akhir Proyek - 3D Simulator & Engine Developer"
    props.author = metadata["author"]
    props.last_modified_by = metadata["author"]
    props.keywords = "Unity WebGL, NavMesh, navigasi spasial, UPNVJ"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main(argv: list[str] | None = None) -> None:
    argv = argv or sys.argv[1:]
    output = Path(argv[0]).resolve() if argv else DEFAULT_OUTPUT
    created = create_base_docx(output)
    print(f"Created clean Faiz base document: {created}")


if __name__ == "__main__":
    main()
